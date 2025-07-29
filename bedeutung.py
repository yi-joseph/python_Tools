import docx
from pypinyin import lazy_pinyin
import requests
from bs4 import BeautifulSoup


def get_wikipedia_definition(word):
    """中文维基百科抓取"""
    url = f"https://zh.wikipedia.org/wiki/{word}"
    try:
        r = requests.get(url, timeout=5)
        r.encoding = 'utf-8'
        soup = BeautifulSoup(r.text, 'html.parser')
        p = soup.find("p")
        if p:
            return p.get_text().strip()
    except:
        return "（未找到释义）"
    return "（未找到释义）"

# ---- Word-Datei einlesen ----
input_file = "woerter.docx"
output_file = "woerter_mit_bedeutung.docx"
doc = docx.Document(input_file)
new_doc = docx.Document()

for para in doc.paragraphs:
    word = para.text.strip()
    if not word:
        continue
    pinyin = " ".join(lazy_pinyin(word))
    wiki_def = get_wikipedia_definition(word)
    
    # Neues Format in die Datei schreiben
    new_doc.add_paragraph(f"{word} ({pinyin}) –")
    new_doc.add_paragraph(f"(Wikipedia)\n{wiki_def}")
    new_doc.add_paragraph("")  # Leerzeile

new_doc.save(output_file)
print(f"Neue Datei gespeichert: {output_file}")
